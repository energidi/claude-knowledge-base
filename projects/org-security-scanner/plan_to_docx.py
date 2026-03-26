from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

PLAN_PATH = r"C:\Users\GidiAbramovich\.claude\plans\scalable-strolling-kahan.md"
OUT_PATH  = r"C:\Users\GidiAbramovich\Documents\Visual Studio Code\Salesforce Security App\Org Security Scanner - Technical Plan.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_run_with_inline(para, text):
    """Add a paragraph run, converting **bold** and `code` inline."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(2):  # **bold**
            r = para.add_run(m.group(2))
            r.bold = True
        elif m.group(3):  # `code`
            r = para.add_run(m.group(3))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])

def style_header_row(table):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, '1F3864')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)

def add_table_from_md(doc, lines):
    # Filter separator rows and parse
    rows = []
    for line in lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_inline(para, cell_text)
            for run in para.runs:
                run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    style_header_row(table)
    doc.add_paragraph()

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        p.paragraph_format.left_indent = Cm(1)
        # light grey shade via direct XML
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)
    doc.add_paragraph()

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Heading styles
for lvl, sz, bold, color_hex in [
    (1, 18, True,  '1F3864'),
    (2, 14, True,  '2E74B5'),
    (3, 11, True,  '2E74B5'),
    (4, 10, True,  '404040'),
]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.size = Pt(sz)
    hs.font.bold = bold
    hs.font.color.rgb = RGBColor(
        int(color_hex[0:2],16),
        int(color_hex[2:4],16),
        int(color_hex[4:6],16)
    )
    hs.font.name = 'Calibri'

# ── parse and render ──────────────────────────────────────────────────────────

with open(PLAN_PATH, encoding='utf-8') as f:
    raw_lines = f.read().splitlines()

i = 0
while i < len(raw_lines):
    line = raw_lines[i]
    stripped = line.strip()

    # Headings
    if stripped.startswith('#### '):
        p = doc.add_heading(stripped[5:], level=4)
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)

    # Horizontal rule
    elif stripped == '---':
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E74B5')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Code block
    elif stripped.startswith('```'):
        i += 1
        code_lines = []
        while i < len(raw_lines) and not raw_lines[i].strip().startswith('```'):
            code_lines.append(raw_lines[i])
            i += 1
        add_code_block(doc, code_lines)

    # Table
    elif stripped.startswith('|'):
        table_lines = []
        while i < len(raw_lines) and raw_lines[i].strip().startswith('|'):
            table_lines.append(raw_lines[i])
            i += 1
        add_table_from_md(doc, table_lines)
        continue  # i already advanced

    # Bullet list
    elif stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_run_with_inline(p, stripped[2:])
        for run in p.runs:
            run.font.size = Pt(10)

    # Sub-bullet (4 spaces or tab + -)
    elif re.match(r'^( {4}|\t)- ', line):
        text = re.sub(r'^( {4}|\t)- ', '', line)
        p = doc.add_paragraph(style='List Bullet 2')
        add_run_with_inline(p, text)
        for run in p.runs:
            run.font.size = Pt(9.5)

    # Numbered list
    elif re.match(r'^\d+\. ', stripped):
        text = re.sub(r'^\d+\. ', '', stripped)
        p = doc.add_paragraph(style='List Number')
        add_run_with_inline(p, text)
        for run in p.runs:
            run.font.size = Pt(10)

    # Empty line
    elif stripped == '':
        pass  # skip blank lines (spacing handled by paragraph spacing)

    # Normal paragraph
    else:
        p = doc.add_paragraph()
        add_run_with_inline(p, stripped)
        for run in p.runs:
            run.font.size = Pt(10)

    i += 1

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
